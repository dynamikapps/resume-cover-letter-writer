import io
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import html


def download_word(content, filename):
    doc = Document()

    # Define styles for section headings and paragraphs
    section_heading_style = doc.styles.add_style(
        'Section Heading', WD_STYLE_TYPE.PARAGRAPH)
    section_heading_style.font.name = 'Arial'
    section_heading_style.font.size = Pt(14)
    section_heading_style.font.bold = True

    paragraph_style = doc.styles.add_style(
        'Paragraph', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_style.font.name = 'Arial'
    paragraph_style.font.size = Pt(12)

    # Split the content into sections
    sections = content.split('\n\n')

    for section in sections:
        lines = section.split('\n')
        if len(lines) > 0:
            # Add section heading
            section_heading = lines[0].strip()
            doc.add_paragraph(section_heading, style='Section Heading')

            # Add section content
            for line in lines[1:]:
                if line.strip():
                    doc.add_paragraph(line.strip(), style='Paragraph')

    # Save the document to a BytesIO object
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    # Provide the download functionality using Streamlit's download_button
    st.download_button(
        label="Download Word Document",
        data=byte_io,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def download_pdf(content, filename):
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        name='HeadingStyle',
        parent=styles['Heading1'],
        fontSize=16,
        leading=24,
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    subheading_style = ParagraphStyle(
        name='SubheadingStyle',
        parent=styles['Heading2'],
        fontSize=14,
        leading=20,
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    custom_style = ParagraphStyle(
        name='CustomStyle',
        parent=styles['Normal'],
        fontSize=12,
        leading=18,
        spaceAfter=12
    )
    elements = []
    sections = content.split('\n\n')
    for section in sections:
        lines = section.split('\n')
        if len(lines) > 0:
            # Add section heading
            section_heading = lines[0].strip()
            elements.append(Paragraph(section_heading, heading_style))

            # Add section subheadings and content
            for line in lines[1:]:
                if line.startswith('- '):
                    elements.append(
                        Paragraph(line[2:].strip(), subheading_style))
                elif line.strip():
                    elements.append(Paragraph(line.strip(), custom_style))

    def canvasmaker(filename):
        doc = SimpleDocTemplate(filename, pagesize=letter, rightMargin=72,
                                leftMargin=72, topMargin=72, bottomMargin=18)
        return doc

    byte_io = io.BytesIO()
    doc = canvasmaker(byte_io)
    doc.build(elements)
    byte_io.seek(0)

    st.download_button(
        label="Download PDF",
        data=byte_io,
        file_name=filename,
        mime="application/pdf"
    )


def display_html_description(description):
    cleaned_description = html.unescape(description)
    st.write(cleaned_description, unsafe_allow_html=True)
